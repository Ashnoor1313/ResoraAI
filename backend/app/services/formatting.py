import fitz  # PyMuPDF
import docx
import io
from typing import Dict, List, Any

# Standard safe ATS fonts
SAFE_FONTS = {
    "arial", "calibri", "times", "times new roman", "garamond", "georgia", 
    "cambria", "helvetica", "tahoma", "trebuchet", "verdana", "courier"
}

def analyze_pdf_formatting(file_bytes: bytes) -> Dict[str, Any]:
    """Analyzes a PDF for ATS visual formatting risks using PyMuPDF."""
    warnings = []
    suggestions = []
    strengths = []
    
    multiple_columns = False
    has_images = False
    has_tables = False
    unsupported_fonts = set()
    has_header_footer_text = False
    total_pages = 0
    
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        total_pages = len(doc)
        
        # Limit processing if resume is extremely long
        for page_num in range(min(total_pages, 5)):
            page = doc[page_num]
            rect = page.rect
            width, height = rect.width, rect.height
            
            # 1. Detect Images
            image_list = page.get_images()
            if image_list:
                has_images = True
            
            # 2. Detect Tables (via PyMuPDF page.find_tables() or vector paths)
            try:
                tables = page.find_tables()
                if tables and len(tables.tables) > 0:
                    has_tables = True
            except Exception:
                # Fallback check on vector drawings representing grids
                drawings = page.get_drawings()
                if len(drawings) > 10:  # Many lines/drawings could indicate a table or graphics
                    has_tables = True
                    
            # 3. Detect Multiple Columns (horizontal blocks overlap)
            # We fetch text blocks with coordinates: (x0, y0, x1, y1, "text", block_no, block_type)
            blocks = page.get_text("blocks")
            text_blocks = [b for b in blocks if b[6] == 0] # Filter only text blocks
            
            # We divide the page width into three regions: left, middle, right.
            # If we find blocks that are side-by-side (overlap vertically but occupy different X ranges),
            # it indicates a multi-column layout.
            side_by_side_count = 0
            for i in range(len(text_blocks)):
                for j in range(i + 1, len(text_blocks)):
                    bi = text_blocks[i]
                    bj = text_blocks[j]
                    
                    # Check vertical overlap
                    y_overlap = max(0, min(bi[3], bj[3]) - max(bi[1], bj[1]))
                    # Check if they are horizontal neighbors
                    bi_width = bi[2] - bi[0]
                    bj_width = bj[2] - bj[0]
                    
                    if y_overlap > 15 and bi_width < width * 0.7 and bj_width < width * 0.7:
                        # Check horizontal separation
                        if (bi[2] < bj[0]) or (bj[2] < bi[0]):
                            side_by_side_count += 1
            
            if side_by_side_count > 3:
                multiple_columns = True
                
            # 4. Check for headers and footers (text very close to top/bottom edges)
            for b in text_blocks:
                y0, y1 = b[1], b[3]
                if y0 < 45 or y1 > height - 45:
                    has_header_footer_text = True
                    
            # 5. Font Analysis
            # Retrieve fonts list on the page
            fonts = page.get_fonts()
            for f in fonts:
                # Font info tuple: (xref, ext, type, fontname, fullname, ...)
                font_name = f[3].lower()
                # Clean prefix from font subset (e.g. "AAAAAA+Calibri" -> "calibri")
                if "+" in font_name:
                    font_name = font_name.split("+")[1]
                
                # Check if it is a safe font
                is_safe = False
                for safe in SAFE_FONTS:
                    if safe in font_name:
                        is_safe = True
                        break
                if not is_safe:
                    unsupported_fonts.add(f[3])
                    
        doc.close()
    except Exception as e:
        warnings.append(f"Could not complete structural layout scan: {str(e)}")
        
    # Compile warnings, suggestions and strengths
    if multiple_columns:
        warnings.append("Multi-column layout detected.")
        suggestions.append("Convert your resume to a single-column layout. ATS parsers read from left-to-right, and side-by-side text columns often get merged or scrambled.")
    else:
        strengths.append("Single-column structure matches standard ATS layout expectations.")
        
    if has_images:
        warnings.append("Embedded images or graphics detected.")
        suggestions.append("Remove profile pictures, icons, charts, or graphical skill bars. ATS scanners cannot read images and they may crash the parsing process.")
    else:
        strengths.append("No profile photos or heavy graphics found, keeping parsing clean.")
        
    if has_tables:
        warnings.append("Tables or complex grids detected.")
        suggestions.append("Avoid wrapping experience or contact info inside tables. Some ATS systems ignore table contents entirely or fail to parse cells in the correct order.")
    else:
        strengths.append("No complex tables detected, preventing cell extraction issues.")
        
    if has_header_footer_text:
        warnings.append("Text detected in header/footer margins.")
        suggestions.append("Move critical contact details out of the top/bottom page margin areas. Many ATS parsers ignore header/footer spaces to avoid parsing page numbers, which means your contact info could be lost.")
    else:
        strengths.append("All key details are kept within the main page boundaries.")
        
    if unsupported_fonts:
        clean_fonts = [f.split("+")[-1] for f in unsupported_fonts]
        warnings.append(f"Non-standard fonts detected: {', '.join(clean_fonts[:3])}.")
        suggestions.append("Stick to standard professional system fonts (e.g., Arial, Calibri, Times New Roman). Custom or decorative fonts can fail to encode text correctly, causing the parser to see garbled characters.")
    else:
        strengths.append("Uses safe, highly compatible standard system typography.")
        
    if total_pages > 2:
        warnings.append(f"Resume length is {total_pages} pages.")
        suggestions.append("Condense your resume to 1 page (if you have under 5 years of experience) or maximum 2 pages. Longer resumes suffer lower recruiter review times.")
    elif total_pages == 1 or total_pages == 2:
        strengths.append(f"Length is optimal ({total_pages} page{'s' if total_pages > 1 else ''}).")
        
    return {
        "multiple_columns": multiple_columns,
        "has_images": has_images,
        "has_tables": has_tables,
        "unsupported_fonts": list(unsupported_fonts),
        "has_header_footer_text": has_header_footer_text,
        "warnings": warnings,
        "suggestions": suggestions,
        "strengths": strengths
    }

def analyze_docx_formatting(file_bytes: bytes) -> Dict[str, Any]:
    """Analyzes a DOCX for ATS visual formatting risks using python-docx."""
    warnings = []
    suggestions = []
    strengths = []
    
    has_tables = False
    has_images = False
    
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        
        # 1. Detect Tables
        if len(doc.tables) > 0:
            has_tables = True
            
        # 2. Detect Images (checking inline shapes and drawings)
        # docx inline_shapes contains pictures, charts
        if len(doc.inline_shapes) > 0:
            has_images = True
            
    except Exception as e:
        warnings.append(f"Could not complete DOCX visual scan: {str(e)}")
        
    if has_tables:
        warnings.append("Tables detected in Word document.")
        suggestions.append("Remove tables and write text in plain paragraphs. ATS parsers can read rows out of sequence.")
    else:
        strengths.append("No table structures found in document.")
        
    if has_images:
        warnings.append("Embedded graphics or shapes detected.")
        suggestions.append("Remove icons, shapes, and pictures to ensure text matches standard parsing guidelines.")
    else:
        strengths.append("No embedded shapes or photos discovered.")
        
    return {
        "multiple_columns": False, # DOCX text is parsed linearly, column layouts are rare parsing issues compared to PDF
        "has_images": has_images,
        "has_tables": has_tables,
        "unsupported_fonts": [],
        "has_header_footer_text": False,
        "warnings": warnings,
        "suggestions": suggestions,
        "strengths": strengths
    }

def analyze_formatting(file_bytes: bytes, filename: str) -> Dict[str, Any]:
    """Main formatting checker selector."""
    if filename.lower().endswith(".pdf"):
        return analyze_pdf_formatting(file_bytes)
    elif filename.lower().endswith(".docx"):
        return analyze_docx_formatting(file_bytes)
    else:
        # Defaults for other formats
        return {
            "multiple_columns": False,
            "has_images": False,
            "has_tables": False,
            "unsupported_fonts": [],
            "has_header_footer_text": False,
            "warnings": [],
            "suggestions": ["Upload your resume in PDF or DOCX format for a thorough layout check."],
            "strengths": ["Document uploaded successfully."]
        }
