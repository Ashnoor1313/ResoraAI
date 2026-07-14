import re
import io
import fitz  # PyMuPDF
import docx
from typing import Dict, List, Any

# Common section header keywords (case-insensitive regex patterns)
SECTION_HEADERS = {
    "education": [r"\beducation\b", r"\bacademic\b", r"\bqualifications\b"],
    "experience": [r"\bexperience\b", r"\bemployment\b", r"\bwork history\b", r"\bprofessional history\b"],
    "projects": [r"\bprojects\b", r"\bacademic projects\b", r"\bpersonal projects\b"],
    "skills": [r"\bskills\b", r"\btechnical skills\b", r"\btechnologies\b", r"\bcompetencies\b"],
    "certifications": [r"\bcertifications\b", r"\bcertificates\b", r"\blicenses\b", r"\badditional activities\b", r"\bawards\b"]
}

# A comprehensive local skill dictionary containing Tech and Non-Tech keywords
SKILLS_DICTIONARY = {
    # Languages
    "python", "javascript", "typescript", "java", "c\\+\\+", "c#", "ruby", "go", "rust", "php", "swift", "kotlin", "sql", "html", "css", "bash", "r", "scalar", "matlab",
    # Frameworks & Libraries
    "react", "angular", "vue", "next\\.js", "nuxt", "svelte", "django", "flask", "fastapi", "spring", "express", "nodejs", "bootstrap", "tailwind", "pandas", "numpy", "scikit-learn", "tensorflow", "pytorch", "keras", "opencv", "graphql", "redux", "jquery", "flutter", "react native",
    # Tools, Platforms & DevOps
    "git", "github", "docker", "kubernetes", "aws", "gcp", "azure", "jenkins", "terraform", "ansible", "linux", "unix", "heroku", "netlify", "vercel", "jira", "confluence",
    # Databases
    "postgresql", "mysql", "mongodb", "sqlite", "redis", "elasticsearch", "cassandra", "mariadb", "dynamodb", "supabase", "firebase",
    # Non-Tech / Business / Design
    "product management", "project management", "scrum", "agile", "seo", "sem", "copywriting", "digital marketing", "marketing strategy", "financial analysis", "excel", "powerpoint", "tableau", "power bi", "figma", "sketch", "photoshop", "illustrator", "wireframing", "ui/ux", "communication", "leadership", "negotiation", "sales", "customer service", "public speaking", "content writing", "recruiting", "talent acquisition", "human resources", "hr", "payroll", "budgeting", "risk management", "strategic planning"
}

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extracts text from PDF using PyMuPDF."""
    text = ""
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        for page in doc:
            text += page.get_text() + "\n"
        doc.close()
    except Exception as e:
        print(f"Error parsing PDF: {e}")
    return text

def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extracts text from DOCX using python-docx."""
    text = ""
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        # Extract paragraph text
        paragraphs = [p.text for p in doc.paragraphs]
        text += "\n".join(paragraphs) + "\n"
        
        # Extract table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\n"
    except Exception as e:
        print(f"Error parsing DOCX: {e}")
    return text

def extract_links_from_pdf_bytes(file_bytes: bytes) -> List[str]:
    """Extracts embedded hyperlinks from PDF using PyMuPDF."""
    links = []
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        for page in doc:
            for link in page.get_links():
                uri = link.get("uri")
                if uri:
                    links.append(uri)
        doc.close()
    except Exception as e:
        print(f"Error extracting links from PDF: {e}")
    return list(set(links))

def extract_links_from_docx_bytes(file_bytes: bytes) -> List[str]:
    """Extracts embedded hyperlinks from DOCX using python-docx."""
    links = []
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        for rel_id, rel in doc.part.rels.items():
            if rel.reltype and "hyperlink" in rel.reltype.lower():
                target = rel.target_ref
                if target and (target.startswith("http://") or target.startswith("https://") or target.startswith("mailto:")):
                    links.append(target)
    except Exception as e:
        print(f"Error extracting links from DOCX: {e}")
    return list(set(links))

def clean_text(text: str) -> str:
    """Cleans extracted text by normalizing whitespaces."""
    # Replace non-breaking spaces, clean tabs
    text = text.replace('\xa0', ' ')
    # Normalize line breaks
    text = re.sub(r'\n+', '\n', text)
    # Remove leading/trailing whitespaces per line
    lines = [line.strip() for line in text.split('\n')]
    return "\n".join([l for l in lines if l])

def parse_contact_info(text: str) -> Dict[str, str]:
    """Helper to parse Name, Email, and Phone using regex and heuristics."""
    email_pattern = r'[\w\.-]+@[\w\.-]+\.\w+'
    # General phone number pattern supporting various formats
    phone_pattern = r'(?:\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
    
    email_match = re.search(email_pattern, text)
    phone_match = re.search(phone_pattern, text)
    
    email = email_match.group(0) if email_match else ""
    phone = phone_match.group(0) if phone_match else ""
    
    # Heuristic for Name: Typically the first line or two of the resume,
    # before contact details, and doesn't contain contact keywords or formatting tokens.
    name = ""
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    
    for line in lines[:5]:
        # Exclude line if it looks like email, phone, or standard resume junk
        if email and email in line:
            continue
        if phone and re.sub(r'[-.\s\(\)+]', '', phone) in re.sub(r'[-.\s\(\)+]', '', line):
            continue
        if any(kw in line.lower() for kw in ["resume", "cv", "curriculum", "page", "portfolio", "github", "linkedin", "address", "street", "http", "www"]):
            continue
        # Names are usually 2-4 words, capitalized
        words = line.split()
        if 1 <= len(words) <= 4 and all(w[0].isupper() or w[0].isdigit() or w in ["de", "von", "van", "di"] for w in words if w):
            name = line
            break
            
    # Fallback to first line if no valid match found
    if not name and lines:
        name = lines[0]
        
    return {"name": name, "email": email, "phone": phone}

def parse_sections(text: str) -> Dict[str, List[str]]:
    """Segments the resume text into standard sections based on keywords."""
    lines = text.split('\n')
    sections = {
        "education": [],
        "experience": [],
        "projects": [],
        "skills": [],
        "certifications": []
    }
    
    current_section = None
    
    for line in lines:
        line_clean = line.strip()
        if not line_clean:
            continue
            
        # Check if line matches any section header
        found_header = False
        # Limit search to short lines (usually headers are < 5 words)
        if len(line_clean.split()) < 5:
            for sec_name, patterns in SECTION_HEADERS.items():
                if any(re.search(pat, line_clean.lower()) for pat in patterns):
                    current_section = sec_name
                    found_header = True
                    break
        
        if found_header:
            continue
            
        if current_section:
            sections[current_section].append(line_clean)
            
    return sections

def extract_skills_from_text(text: str) -> List[str]:
    """Scans text against a preset skill dictionary to extract matches."""
    text_lower = text.lower()
    detected_skills = set()
    
    for skill in SKILLS_DICTIONARY:
        # Match using word boundaries. Add support for specialized signs in programming languages (e.g. C++, C#)
        pattern = r'\b' + skill + r'\b'
        if skill == "c\\+\\+":
            pattern = r'c\+\+'
        elif skill == "c#":
            pattern = r'c#'
        elif skill == "next\\.js":
            pattern = r'next\.js|nextjs'
            
        if re.search(pattern, text_lower):
            # Clean skill presentation
            clean_skill_name = skill.replace('\\', '')
            detected_skills.add(clean_skill_name.title() if clean_skill_name not in ["aws", "gcp", "hr", "seo", "sem", "sql", "html", "css", "ui/ux"] else clean_skill_name.upper())
            
    return sorted(list(detected_skills))

def parse_resume(file_bytes: bytes, filename: str) -> Dict[str, Any]:
    """Main parsing orchestrator that returns structured resume content."""
    # 1. Text extraction
    links = []
    if filename.lower().endswith(".pdf"):
        raw_text = extract_text_from_pdf(file_bytes)
        links = extract_links_from_pdf_bytes(file_bytes)
    elif filename.lower().endswith(".docx"):
        raw_text = extract_text_from_docx(file_bytes)
        links = extract_links_from_docx_bytes(file_bytes)
    else:
        # Fallback to plain text
        try:
            raw_text = file_bytes.decode("utf-8")
        except Exception:
            raw_text = str(file_bytes)
            
    clean_raw_text = clean_text(raw_text)
    
    # 2. Extract metadata
    contact = parse_contact_info(clean_raw_text)
    sections = parse_sections(clean_raw_text)
    skills = extract_skills_from_text(clean_raw_text)
    
    return {
        "name": contact["name"],
        "email": contact["email"],
        "phone": contact["phone"],
        "skills": skills,
        "education": sections["education"],
        "experience": sections["experience"],
        "projects": sections["projects"],
        "certifications": sections["certifications"],
        "raw_text": clean_raw_text,
        "links": links
    }
